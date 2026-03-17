"""CV parser: reads PDF / DOCX / TXT files and extracts a structured research profile.

The structured profile (CVProfile) is research-oriented: it captures thesis topics,
publications, research interests and methodology — fields that matter for PhD/postdoc
applications and are typically absent from industry-focused parsers.
"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, TypedDict

from agent.llm_client import LLMClient


# ---------------------------------------------------------------------------
# Type definitions
# ---------------------------------------------------------------------------

class CVProfile(TypedDict, total=False):
    name: str
    contact: dict                  # email, phone, linkedin, github, website
    summary: str
    education: list[dict]          # degree, institution, field, year, thesis_topic
    experience: list[dict]         # title, institution/company, dates, description
    research_interests: list[str]
    publications: list[dict]       # title, venue, year, authors
    skills: dict                   # programming, tools, languages, lab_techniques
    awards: list[str]
    languages: list[dict]          # language, level
    references: list[dict]         # name, title, institution


# ---------------------------------------------------------------------------
# Raw text extraction helpers
# ---------------------------------------------------------------------------

def _extract_text_from_pdf(path: Path) -> str:
    """Extract all text from a PDF using pdfplumber."""
    try:
        import pdfplumber  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "pdfplumber is required to parse PDF files.\n"
            "Install it with:  pip install pdfplumber"
        ) from exc

    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n".join(pages)


def _extract_text_from_docx(path: Path) -> str:
    """Extract all text from a DOCX using python-docx."""
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "python-docx is required to parse DOCX files.\n"
            "Install it with:  pip install python-docx"
        ) from exc

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also grab text from tables (common in European CVs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def _extract_text_from_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def extract_raw_text(cv_path: str | Path) -> str:
    """Extract raw text from a CV file.

    Supported formats: .pdf, .docx, .doc, .txt

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file extension is unsupported.
    """
    path = Path(cv_path)
    if not path.exists():
        raise FileNotFoundError(f"CV file not found: {path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_text_from_pdf(path)
    elif suffix in (".docx", ".doc"):
        return _extract_text_from_docx(path)
    elif suffix == ".txt":
        return _extract_text_from_txt(path)
    else:
        raise ValueError(
            f"Unsupported CV format '{suffix}'. "
            "Supported formats: .pdf, .docx, .doc, .txt"
        )


# ---------------------------------------------------------------------------
# LLM-powered structured extraction
# ---------------------------------------------------------------------------

_EXTRACTION_SYSTEM = (
    "You are an expert academic CV parser. "
    "Extract ALL information from the CV text into structured JSON. "
    "Pay special attention to research-specific fields: thesis topics, "
    "publications, research interests, methodologies, and lab/technical skills. "
    "Respond only with valid JSON — no markdown fences, no commentary."
)

_EXTRACTION_PROMPT = """Extract ALL information from the following CV and return a single JSON object.

Expected structure (use null for missing scalar fields, [] for missing lists):
{{
  "name": "Full Name",
  "contact": {{
    "email": "email@example.com",
    "phone": "+39 000 0000000",
    "linkedin": "linkedin.com/in/...",
    "github": "github.com/...",
    "website": "https://..."
  }},
  "summary": "Brief professional/research summary",
  "education": [
    {{
      "degree": "PhD",
      "institution": "MIT",
      "field": "Computer Science",
      "year": "2021",
      "thesis_topic": "Deep reinforcement learning for robotic control"
    }}
  ],
  "experience": [
    {{
      "title": "Research Assistant",
      "institution": "ETH Zurich",
      "dates": "2019-2021",
      "description": "Worked on NLP models for low-resource languages."
    }}
  ],
  "research_interests": ["machine learning", "natural language processing"],
  "publications": [
    {{
      "title": "Paper Title",
      "venue": "NeurIPS 2022",
      "year": "2022",
      "authors": "Smith J., Doe A."
    }}
  ],
  "skills": {{
    "programming": ["Python", "C++", "R"],
    "tools": ["PyTorch", "TensorFlow", "Docker"],
    "languages": ["LaTeX", "MATLAB"],
    "lab_techniques": ["PCR", "cell culture"]
  }},
  "awards": ["Best Paper Award NeurIPS 2022", "EPSRC Scholarship"],
  "languages": [
    {{"language": "English", "level": "Native"}},
    {{"language": "Italian", "level": "B2"}}
  ],
  "references": [
    {{"name": "Prof. Jane Smith", "title": "Full Professor", "institution": "MIT"}}
  ]
}}

Do NOT invent information — extract only what is present.

CV TEXT:
---
{cv_text}
---"""


def _strip_fences(text: str) -> str:
    """Remove markdown code fences that some models emit."""
    text = re.sub(r"^```(?:json)?\s*", "", text.strip())
    text = re.sub(r"\s*```$", "", text)
    return text.strip()


def extract_profile(
    cv_path: str | Path,
    model: str | None = None,
    console=None,
) -> CVProfile:
    """Parse a CV file and return a structured CVProfile dict.

    Args:
        cv_path: Path to the CV file (.pdf, .docx, .txt).
        model:   Optional LLM model override.
        console: Optional rich Console for progress messages.

    Returns:
        A CVProfile dictionary with all extracted fields.
    """
    if console:
        console.print("[cyan]Extracting raw text from CV...[/cyan]")

    raw_text = extract_raw_text(cv_path)
    if not raw_text.strip():
        raise ValueError("Could not extract any text from the CV file.")

    if console:
        console.print(
            f"[green]Extracted {len(raw_text):,} characters.[/green] "
            "Sending to LLM for structured extraction..."
        )

    client = LLMClient(model=model)
    prompt = _EXTRACTION_PROMPT.format(cv_text=raw_text[:8000])  # truncate for context window

    raw_json = client.generate(
        system=_EXTRACTION_SYSTEM,
        user=prompt,
        json_mode=True,
    )

    raw_json = _strip_fences(raw_json)

    try:
        profile: CVProfile = json.loads(raw_json)
    except json.JSONDecodeError:
        # Try to salvage a JSON object from the response
        m = re.search(r"\{.*\}", raw_json, re.DOTALL)
        if m:
            try:
                profile = json.loads(m.group())
            except json.JSONDecodeError:
                profile = {"name": "Unknown", "summary": raw_json[:500]}
        else:
            profile = {"name": "Unknown", "summary": raw_json[:500]}

    return profile


# ---------------------------------------------------------------------------
# Profile → compact text summary (for use in downstream prompts)
# ---------------------------------------------------------------------------

def profile_summary(profile: CVProfile) -> str:
    """Build a compact text summary of a CVProfile for use in LLM prompts."""
    lines: list[str] = []

    if profile.get("name"):
        lines.append(f"Name: {profile['name']}")

    contact: dict = profile.get("contact") or {}
    if contact.get("email"):
        lines.append(f"Email: {contact['email']}")

    if profile.get("summary"):
        lines.append(f"Summary: {profile['summary']}")

    research = profile.get("research_interests") or []
    if research:
        lines.append(f"Research interests: {', '.join(research[:10])}")

    edu = profile.get("education") or []
    if edu:
        for e in edu[:3]:
            thesis = f" — Thesis: {e['thesis_topic']}" if e.get("thesis_topic") else ""
            lines.append(
                f"Education: {e.get('degree', '')} in {e.get('field', '')} "
                f"from {e.get('institution', '')} ({e.get('year', '')}){thesis}"
            )

    pubs = profile.get("publications") or []
    if pubs:
        lines.append(f"Publications ({len(pubs)}):")
        for p in pubs[:5]:
            lines.append(f"  - \"{p.get('title', '')}\" — {p.get('venue', '')} {p.get('year', '')}")

    exp = profile.get("experience") or []
    if exp:
        lines.append(f"Experience ({len(exp)} positions):")
        for e in exp[:4]:
            lines.append(
                f"  - {e.get('title', '')} at {e.get('institution', '')} ({e.get('dates', '')})"
            )

    skills: dict = profile.get("skills") or {}
    prog = skills.get("programming") or []
    tools = skills.get("tools") or []
    all_skills = prog + tools
    if all_skills:
        lines.append(f"Technical skills: {', '.join(all_skills[:20])}")
    lab = skills.get("lab_techniques") or []
    if lab:
        lines.append(f"Lab techniques: {', '.join(lab[:10])}")

    awards = profile.get("awards") or []
    if awards:
        lines.append(f"Awards: {'; '.join(awards[:5])}")

    langs = profile.get("languages") or []
    if langs:
        lang_str = ", ".join(
            f"{la.get('language', '')} ({la.get('level', '')})" for la in langs
        )
        lines.append(f"Languages: {lang_str}")

    return "\n".join(lines)
